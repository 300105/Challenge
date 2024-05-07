# Importando as bibliotecas necessárias
import os
import sys
from PIL import Image, ImageTk
from docx import Document
from openpyxl import load_workbook
from subprocess import run
from tkinter import *
from tkinter import messagebox, ttk, filedialog, Text
import openpyxl
import PyPDF2
import tkinter as tk
import DataBaser
from PIL import Image, ImageTk
import cv2
import pytesseract

# Classe para redirecionar a saída padrão para um widget de texto
class TextRedirector(object):
    def __init__(self, widget):
        self.widget = widget

    def write(self, str):
        self.widget.insert("end", str)
        self.widget.see("end")

# Função para redirecionar a saída padrão para um widget de texto
def sua_funcao():
    global output
    output = Text(RightFrame, width=38, height=10)
    output.place(x=40, y=120)
    sys.stdout = TextRedirector(output)

# Função para procurar uma palavra em um arquivo txt
def procurar_palavra_em_txt(txt_file, palavra):
    with open(txt_file, 'r', encoding='utf-8') as file:
        texto = file.read()
        if palavra.lower() in texto.lower():
            return True
    return False

# Função para procurar uma palavra em um arquivo docx
def procurar_palavra_em_docx(docx_file, palavra):
    doc = Document(docx_file)
    for para in doc.paragraphs:
        if palavra.lower() in para.text.lower():
            print("A palavra '{}' foi encontrada no documento Word.".format(palavra))
            return

# Função para procurar uma palavra em um arquivo excel
def procurar_palavra_em_excel(excel_file, palavra):
    wb = openpyxl.load_workbook(excel_file)
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        for row in sheet.iter_rows():
            for cell in row:
                if palavra.lower() in str(cell.value).lower():
                    print("A palavra '{}' foi encontrada na planilha '{}'.".format(palavra, sheet_name))
                    return
                
def procurar_palavras_em_imagem(imagem, palavras):
    # Carregar a imagem
    imagem = cv2.imread(imagem)
    
    # Usar o Tesseract para fazer a OCR (Reconhecimento Óptico de Caracteres)
    texto = pytesseract.image_to_string(imagem)
    
    # Procurar as palavras no texto extraído
    palavras_encontradas = []
    for palavra in palavras:
        if palavra.lower() in texto.lower():
            palavras_encontradas.append(palavra)
            #print("A palavra '{}' foi encontrada na imagem.".format(palavra))

    if not palavras_encontradas:
        print("Nenhuma das palavras foi encontrada na imagem.")
    return palavras_encontradas

# Função para selecionar uma pasta
def selecionar_pasta():
    for widget in RightFrame.winfo_children():
        widget.destroy()
    SelecionarPasta = ttk.Button(RightFrame, text="Selecione sua Pasta", width=50, command=procurar_arquivos_em_pasta)
    SelecionarPasta.place(x=40, y=80)
    output = Text(RightFrame, width=38, height=10)
    output.place(x=40, y=120)
    sys.stdout = TextRedirector(output)

# Função para procurar arquivos em uma pasta
def procurar_arquivos_em_pasta():
    dir_path = filedialog.askdirectory()
    if dir_path:
        palavras = ["RG", "CPF", "Rua", "Avenida", "Travessa", "Alameda", "Telefone", 
                    "Celular", "@hotmail", "@gmail", "@outllok", "@icloud", "Email:", 
                    "Email", "Católica", "Evangélica", "Espírita", "Umbanda", 
                    "candomblé ", "Ateu", "Judaica", "religiões", "Religião", "Registro geral",
                    "filiação", "naturalidade", "doc original", "carteira de identidade"]
        for filename in os.listdir(dir_path):
            file_path = os.path.join(dir_path, filename)
            if file_path.endswith(".txt") or file_path.endswith(".pdf") or file_path.endswith(".docx") or file_path.endswith(".xlsx"):
                palavras_encontradas = procurar_palavras_em_txt_pdf_docx_xlsx(file_path, palavras)
                if palavras_encontradas:
                    print(f"- A palavra {palavras_encontradas} foi encontrada no arquivo: {file_path}\n")
            elif file_path.endswith(".png"):
                palavras_encontradas = procurar_palavras_em_imagem(file_path, palavras)
                if palavras_encontradas:
                    print(f"- A palavra {palavras_encontradas} foi encontrada na imagem: {file_path}\n")

# Função para procurar palavras em arquivos txt, pdf, docx e xlsx
def procurar_palavra_em_imagem(imagem, palavra):
    # Carregar a imagem
    imagem = cv2.imread(imagem)
    
    # Usar o Tesseract para fazer a OCR (Reconhecimento Óptico de Caracteres)
    texto = pytesseract.image_to_string(imagem)
    
    # Procurar a palavra no texto extraído
    if palavra.lower() in texto.lower():
        print("A palavra '{}' foi encontrada na imagem.".format(palavra))
    else:
        print("A palavra '{}' não foi encontrada na imagem.".format(palavra))

def procurar_palavras_em_txt_pdf_docx_xlsx(file_path, palavras):
    palavras_encontradas = []
    if file_path.endswith(".txt"):
        with open(file_path, 'r', encoding='utf-8') as file:
            texto = file.read()
            for palavra in palavras:
                if palavra.lower() in texto.lower():
                    palavras_encontradas.append(palavra)
    elif file_path.endswith(".pdf"):
        pdf_file_obj = open(file_path, 'rb')
        pdf_reader = PyPDF2.PdfReader(pdf_file_obj)
        num_pages = len(pdf_reader.pages)
        for i in range(num_pages):
            page_obj = pdf_reader.pages[i]
            texto = page_obj.extract_text()
            for palavra in palavras:
                if palavra.lower() in texto.lower():
                    palavras_encontradas.append(palavra)
        pdf_file_obj.close()
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        for para in doc.paragraphs:
            for palavra in palavras:
                if palavra.lower() in para.text.lower():
                    palavras_encontradas.append(palavra)
    elif file_path.endswith(".xlsx"):
        wb = load_workbook(filename = file_path, read_only=True)
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.rows:
                for cell in row:
                    for palavra in palavras:
                        if palavra.lower() in str(cell.value).lower():
                            palavras_encontradas.append(palavra)
    return palavras_encontradas

# Função para selecionar um arquivo
def selecionar_arquivo():
    for widget in RightFrame.winfo_children():
        widget.destroy()
    SelecionarArquivo = ttk.Button(RightFrame, text="Selecione seu Arquivo", width=50, command=procurar_arquivo)
    SelecionarArquivo.place(x=40, y=80)
    output = Text(RightFrame, width=38, height=10)
    output.place(x=40, y=120)
    sys.stdout = TextRedirector(output)

# Função para procurar um arquivo
def procurar_arquivo():
    file_path = filedialog.askopenfilename()
    if file_path:
        palavras = ["RG", "CPF", "Rua", "Avenida", "Travessa", "Alameda", "Telefone", 
                    "Celular", "@hotmail", "@gmail", "@outllok", "@icloud", "Email:", 
                    "Email", "Católica", "Evangélica", "Espírita", "Umbanda", 
                    "candomblé ", "Ateu", "Judaica", "religiões", "Religião", "Registro geral",
                    "filiação", "naturalidade", "doc original", "carteira de identidade"]
        if file_path.endswith(".txt") or file_path.endswith(".pdf") or file_path.endswith(".docx") or file_path.endswith(".xlsx"):
            palavras_encontradas = procurar_palavras_em_txt_pdf_docx_xlsx(file_path, palavras)
            if palavras_encontradas:
                print(f"- As palavras {palavras_encontradas} foram encontradas no arquivo: {file_path}\n")
            else:
                print(f"- Nada foi encontrado no arquivo.\n")
        elif file_path.endswith(".png"):
            palavras_encontradas = procurar_palavras_em_imagem(file_path, palavras)
            if palavras_encontradas:
                print(f"- As palavras {palavras_encontradas} foram encontradas na imagem: {file_path}\n")
            else:
                print(f"- Nada foi encontrado na imagem.\n")

# Função para limpar e adicionar labels
def limpar_e_adicionar_labels():
    for widget in RightFrame.winfo_children():
        widget.destroy()
    Nomes1 = Label(RightFrame, text="Isabelli Barreto                 RM: 553649", font=("Century Gothic", 15), bg="MIDNIGHTBLUE", fg="white")
    Nomes1.place(x=5, y=100)
    Nomes2 = Label(RightFrame, text="Matheus Machado          RM: 553522", font=("Century Gothic", 15), bg="MIDNIGHTBLUE", fg="white")
    Nomes2.place(x=5, y=130)
    Nomes3 = Label(RightFrame, text="Pedro Fernandes             RM: 553501", font=("Century Gothic", 15), bg="MIDNIGHTBLUE", fg="white")
    Nomes3.place(x=5, y=160)
    Nomes4 = Label(RightFrame, text="Pedro Verrone                 RM: 553670", font=("Century Gothic", 15), bg="MIDNIGHTBLUE", fg="white")
    Nomes4.place(x=5, y=190)

# Criando a janela principal
jan = Tk()
jan.title("DP Systems - Acess Panel")
jan.geometry("600x300")
jan.configure(background="white")
jan.resizable(width=False, height=False)
jan.attributes("-alpha", 1.0)
jan.iconbitmap(default=r"C:\\Users\\barre\\OneDrive\\Área de Trabalho\\Entrega 7.05\\Digital defense unit logo.ico")

# Carregando a imagem do logo
image = Image.open(r"C:\\Users\\barre\\OneDrive\\Área de Trabalho\\Entrega 7.05\\Digital defense unit.jpg")
image = image.resize((330, 210), Image.LANCZOS)

logo = ImageTk.PhotoImage(image)

# Criando os frames esquerdo e direito
LeftFrame = Frame(jan, width=200, height=350, bg="MIDNIGHTBLUE", relief="raise")
LeftFrame.pack(side=LEFT)

RightFrame = Frame(jan, width=395, height=300, bg="MIDNIGHTBLUE", relief="raise")
RightFrame.pack(side=RIGHT)

# Criando os botões para ler arquivo, ler pasta e ir para a página inicial
LerArquivo = ttk.Button(LeftFrame, text="Ler Arquivo", width=25, command=selecionar_arquivo)
LerArquivo.place(x=13, y=140)

LerPasta = ttk.Button(LeftFrame, text="Ler Pasta", width=25, command=selecionar_pasta)
LerPasta.place(x=13, y=190)

Home = ttk.Button(LeftFrame, text="Página Inicial", width=25, command=limpar_e_adicionar_labels)
Home.place(x=13, y=90)

# Adicionando o logo ao frame direito
LogoLabel = Label(RightFrame, image=logo, bg="MIDNIGHTBLUE")
LogoLabel.place(x=30, y=45)

# Iniciando o loop principal da janela
jan.mainloop()